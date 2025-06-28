import os
import tempfile
from docx import Document

# --- CONFIGURACIÓN ---
# Asegúrate de que esta ruta sea la misma que en tu script principal
DOWNLOAD_FOLDER = os.path.join(tempfile.gettempdir(), 'informes_descargados_senapred')

def diagnosticar_tablas():
    """
    Abre el último archivo .docx descargado y muestra el contenido de sus tablas.
    """
    try:
        # Encontrar el archivo .docx más reciente en la carpeta de descargas
        files = [os.path.join(DOWNLOAD_FOLDER, f) for f in os.listdir(DOWNLOAD_FOLDER) if f.lower().endswith('.docx')]
        if not files:
            print(f"ERROR: No se encontraron archivos .docx en la carpeta: {DOWNLOAD_FOLDER}")
            return

        latest_file = max(files, key=os.path.getmtime)
        print(f"--- Analizando el archivo: {os.path.basename(latest_file)} ---\n")

        document = Document(latest_file)

        if not document.tables:
            print("El documento no contiene ninguna tabla.")
            return

        print(f"El documento tiene un total de {len(document.tables)} tablas.\n")

        # Iterar sobre cada tabla e imprimir su índice y contenido
        for i, table in enumerate(document.tables):
            print(f"=============== ÍNDICE DE TABLA: {i} ================")
            try:
                # Imprimir las primeras 3 filas como muestra
                for row_idx, row in enumerate(table.rows):
                    if row_idx >= 3: # Limitar la muestra a 3 filas
                        break
                    # Extraer el texto de cada celda de la fila
                    row_content = " | ".join([cell.text.strip() for cell in row.cells])
                    print(f"  Fila {row_idx}: {row_content}")
            except Exception as e:
                print(f"  Error al leer la tabla {i}: {e}")
            print("-" * 50 + "\n")

    except FileNotFoundError:
        print(f"ERROR: La carpeta de descarga no existe: {DOWNLOAD_FOLDER}")
    except Exception as e:
        print(f"Ha ocurrido un error inesperado: {e}")

if __name__ == "__main__":
    diagnosticar_tablas()