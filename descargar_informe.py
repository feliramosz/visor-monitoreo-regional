import imaplib
import email
import os
from dotenv import load_dotenv
load_dotenv()
from datetime import datetime
from docx import Document
import json
import argparse
import re
import tempfile

# --- CONFIGURACIÓN ---
EMAIL_USER = os.getenv('GMAIL_USER', 'monitoreoregionaleco5@gmail.com')
EMAIL_PASSWORD = os.getenv('GMAIL_APP_PASSWORD')
IMAP_SERVER = "imap.gmail.com"
DOWNLOAD_FOLDER = os.path.join(tempfile.gettempdir(), 'informes_descargados_senapred')
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_OUTPUT_FOLDER = os.path.join(SCRIPT_DIR, 'datos_extraidos')
DATA_FILE = os.path.join(DATA_OUTPUT_FOLDER, "ultimo_informe.json")

# --- FUNCIÓN PARA OBTENER UN ID ÚNICO DEL INFORME ---
def obtener_id_del_informe(nombre_archivo):
    """
    Extrae un ID único del nombre del archivo, como 'AM-2025-06-12'.
    Ahora maneja formatos de fecha más flexibles (DD MM YY, DD-MM-YYYY, etc.).
    """
    nombre_archivo_upper = nombre_archivo.upper()
    
    tipo_informe = "AM" if "AM" in nombre_archivo_upper else "PM" if "PM" in nombre_archivo_upper else None
    if not tipo_informe:
        return None

    # Regex : busca DD, MM, y YY o YYYY, separados por espacio, punto o guion.
    match = re.search(r'(\d{1,2})[\s.-]+(\d{1,2})[\s.-]+(\d{2,4})', nombre_archivo)
    if match:
        dia, mes, anio = match.groups()
        
        # Asegurar que día y mes tengan dos dígitos
        dia = dia.zfill(2)
        mes = mes.zfill(2)

        # Convertir año de 2 dígitos a 4 dígitos
        if len(anio) == 2:
            anio = f"20{anio}"
            
        return f"{tipo_informe}-{anio}-{mes}-{dia}"
        
    return None

# --- FUNCIÓN PARA DESCARGAR INFORMES ---
def descargar_ultimo_informe(email_user, email_password, imap_server, download_folder):
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Iniciando descarga del último informe...")

    try:
        mail = imaplib.IMAP4_SSL(imap_server)
        mail.login(email_user, email_password)
        print("Conexión exitosa a Gmail.")
        mail.select('inbox')

        status, email_ids = mail.search(None, 'ALL')
        list_of_emails = email_ids[0].split()

        if not list_of_emails:
            print("No hay correos en la bandeja de entrada.")
            mail.logout()
            return None, None

        latest_email_id = list_of_emails[-1]
        status, msg_data = mail.fetch(latest_email_id, '(RFC822)')
        raw_email = msg_data[0][1]
        msg = email.message_from_bytes(raw_email)

        subject = email.header.decode_header(msg['Subject'])[0][0]
        if isinstance(subject, bytes):
            try:
                subject = subject.decode('utf-8')
            except UnicodeDecodeError:
                subject = subject.decode('latin-1', errors='ignore')

        print(f"Procesando el último correo: '{subject}' de '{msg['from']}'")

        for part in msg.walk():
            if part.get_content_maintype() == 'multipart' or part.get('Content-Disposition') is None:
                continue

            file_name = part.get_filename()

            if file_name and file_name.lower().endswith('.docx'):
                filepath = os.path.join(download_folder, file_name)
                os.makedirs(download_folder, exist_ok=True)

                with open(filepath, 'wb') as f:
                    f.write(part.get_payload(decode=True))
                print(f"Archivo adjunto '{file_name}' descargado en '{download_folder}'")

                mail.store(latest_email_id, '+FLAGS', '\\Seen')
                print(f"Correo marcado como leído: {subject}")

                mail.logout()
                return filepath, subject

        print("El último correo no contiene un archivo .docx adjunto esperado.")
        mail.logout()
        return None, None

    except Exception as e:
        print(f"Error al descargar el informe: {e}")
        return None, None

# --- FUNCIÓN PARA EXTRAER DATOS (MODIFICADA para recibir el ID del informe) ---
def extraer_datos_docx(docx_filepath, subject_email, report_id):
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Iniciando extracción de datos de: {docx_filepath}")
    datos_extraidos = {}

    try:
        # --- Añadir el ID del informe al principio ---
        datos_extraidos['id_informe_origen'] = report_id
        
        document = Document(docx_filepath)
               
        datos_extraidos['tipo_informe'] = 'Desconocido'

        nombre_base_archivo = os.path.basename(docx_filepath).upper()

        if "AM" in nombre_base_archivo:
            datos_extraidos['tipo_informe'] = 'AM'
        elif "PM" in nombre_base_archivo:
            datos_extraidos['tipo_informe'] = 'PM'
                
        datos_extraidos['fecha_informe'] = 'N/A'
        datos_extraidos['hora_informe'] = 'N/A'
        if len(document.tables) > 0:
            table_fecha_hora = document.tables[0]
            try:                
                fecha_str = table_fecha_hora.cell(1, 0).text.strip()
                hora_str = table_fecha_hora.cell(1, 1).text.strip()
                datos_extraidos['fecha_informe'] = fecha_str
                datos_extraidos['hora_informe'] = hora_str
                print(f"Fecha y Hora extraídas: {fecha_str}, {hora_str}")

                if datos_extraidos['tipo_informe'] == 'Desconocido':
                    hora_limpia_str = hora_str.replace('h.', '').strip()
                    try:
                        hora_obj = datetime.strptime(hora_limpia_str, '%H:%M').time()
                        if hora_obj < datetime.strptime('12:00', '%H:%M').time():
                            datos_extraidos['tipo_informe'] = 'AM'
                        else:
                            datos_extraidos['tipo_informe'] = 'PM'
                    except ValueError:
                        print(f"Advertencia: No se pudo parsear la hora '{hora_limpia_str}' para determinar AM/PM.")
            except IndexError:
                print("No se pudo extraer fecha/hora de la primera tabla. Posible cambio de formato.")
        else:
            print("No se encontraron tablas en el documento para fecha/hora.")

        print(f"Tipo de informe identificado: {datos_extraidos['tipo_informe']}")

        # --- EXTRAER RESUMEN DE ALERTAS VIGENTES  ---
        datos_extraidos['alertas_vigentes'] = []
        if len(document.tables) > 1:
            table_alertas = document.tables[1]
            for i, row in enumerate(table_alertas.rows):
                if i == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 4 and any(cells):
                    datos_extraidos['alertas_vigentes'].append({
                        'nivel_alerta': cells[0],
                        'evento': cells[1],
                        'cobertura': cells[2],
                        'amplitud': cells[3]
                    })
            print(f"Alertas vigentes extraídas: {len(datos_extraidos['alertas_vigentes'])} registros.")
        else:
            print("No se encontró la tabla de Alertas Vigentes.")

        # --- EXTRAER RESUMEN DE EMERGENCIAS DE LAS ÚLTIMAS 24 HORAS ---
        datos_extraidos['emergencias_ultimas_24_horas'] = []
        if len(document.tables) > 2:
            table_emergencias = document.tables[2]
            for i, row in enumerate(table_emergencias.rows):
                if i == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 4:
                    if cells[0] and cells[1] and cells[2] and cells[3]:
                        datos_extraidos['emergencias_ultimas_24_horas'].append({
                            'n_informe': cells[0],
                            'fecha_hora': cells[1],
                            'evento_lugar': cells[2],
                            'resumen': cells[3]
                        })
                    elif not cells[0] and not cells[1] and not cells[2] and cells[3]:
                        if datos_extraidos['emergencias_ultimas_24_horas']:
                            datos_extraidos['emergencias_ultimas_24_horas'][-1]['resumen'] += " " + cells[3]

            print(f"Emergencias extraídas: {len(datos_extraidos['emergencias_ultimas_24_horas'])} registros.")
        else:
            print("No se encontró la tabla de Emergencias.")

        # --- EXTRAER AVISOS / ALERTAS / ALARMAS METEOROLÓGICAS ---
        datos_extraidos['avisos_alertas_meteorologicas'] = []
        if len(document.tables) > 4:
            table_avisos_met = document.tables[4]
            for i, row in enumerate(table_avisos_met.rows):
                if i == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 4 and any(cells):
                    datos_extraidos['avisos_alertas_meteorologicas'].append({
                        'aviso_alerta_alarma': cells[0],
                        'fecha_hora_emision': cells[1],
                        'descripcion': cells[2],
                        'cobertura': cells[3]
                    })
            print(f"Avisos/Alertas Meteorológicas extraídas: {len(datos_extraidos['avisos_alertas_meteorologicas'])} registros.")
        else:
            print("No se encontró la tabla de Avisos/Alertas Meteorológicas.")

        # --- EXTRAER 6. ÍNDICE DE RADIACIÓN ULTRAVIOLETA ---
        datos_extraidos['radiacion_uv'] = {
            'observado_ayer_label': 'Observado (sin datos):',
            'observado_ayer_value': 'N/A',
            'pronosticado_hoy_label': 'Pronosticado (sin datos):',
            'pronosticado_hoy_value': 'N/A'
        }
        if len(document.tables) > 12:
            table_uv = document.tables[12]
            try:
                if len(table_uv.rows) >= 2 and len(table_uv.rows[0].cells) >= 2:
                    datos_extraidos['radiacion_uv']['observado_ayer_label'] = table_uv.cell(0, 0).text.strip() + ":"
                    datos_extraidos['radiacion_uv']['pronosticado_hoy_label'] = table_uv.cell(0, 1).text.strip() + ":"
                    datos_extraidos['radiacion_uv']['observado_ayer_value'] = table_uv.cell(1, 0).text.strip()
                    datos_extraidos['radiacion_uv']['pronosticado_hoy_value'] = table_uv.cell(1, 1).text.strip()
                    print(f"Índice UV extraído: {datos_extraidos['radiacion_uv']['observado_ayer_label']} {datos_extraidos['radiacion_uv']['observado_ayer_value']}, {datos_extraidos['radiacion_uv']['pronosticado_hoy_label']} {datos_extraidos['radiacion_uv']['pronosticado_hoy_value']}")
                else:
                    print("Advertencia: Estructura de tabla UV inesperada. No tiene suficientes filas/columnas.")
            except IndexError as e:
                print(f"No se pudo extraer el Índice UV. Error: {e}")
            except Exception as e:
                print(f"Error inesperado al extraer Índice UV: {e}")
        else:
            print("No se encontró la tabla de Índice de Radiación Ultravioleta (o el índice es incorrecto).")

        # --- EXTRAER 8. ESTADO DE CARRETERAS ---
        datos_extraidos['estado_carreteras'] = []
        if len(document.tables) > 14:
            table_carreteras = document.tables[14]
            for i, row in enumerate(table_carreteras.rows):
                if i == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 3 and any(cells):
                    datos_extraidos['estado_carreteras'].append({
                        'carretera': cells[0],
                        'estado': cells[1],
                        'condicion': cells[2]
                    })
            print(f"Estado de Carreteras extraído: {len(datos_extraidos['estado_carreteras'])} registros.")
        else:
            print("No se encontró la tabla de Estado de Carreteras.")

        # --- EXTRAER 9. ESTADO DE PUERTOS ---
        datos_extraidos['estado_puertos'] = []
        if len(document.tables) > 15:
            table_puertos = document.tables[15]
            for i, row in enumerate(table_puertos.rows):
                if i == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 3 and any(cells):
                    datos_extraidos['estado_puertos'].append({
                        'puerto': cells[0],
                        'estado_del_puerto': cells[1],
                        'condicion': cells[2]
                    })
            print(f"Estado de Puertos extraído: {len(datos_extraidos['estado_puertos'])} registros.")
        else:
            print("No se encontró la tabla de Estado de Puertos.")

        # --- EXTRAER 10. DATOS HIDROMÉTRICOS ESTÁTICOS ---
        datos_extraidos['datos_hidrometricos'] = []        
        if len(document.tables) > 13:
            table_hidro = document.tables[13]
            for i, row in enumerate(table_hidro.rows):
                if i == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                                
                if len(cells) >= 4:
                    nombre_estacion = cells[0]                    
                    nivel_caudal_str = cells[2]
                    
                    nivel_m = None
                    caudal_m3s = None

                    try:
                        # Nueva lógica más robusta para manejar el formato "Nivel/Caudal/seg"
                        parts = nivel_caudal_str.split('/')

                        # Nos aseguramos de que haya al menos 2 partes (Nivel y Caudal)
                        if len(parts) >= 2:
                            # Extraemos el número de la primera parte (Nivel)
                            nivel_str_match = re.search(r'[\d\.]+', parts[0])
                            if nivel_str_match:
                                nivel_m = float(nivel_str_match.group(0))

                            # Extraemos el número de la segunda parte (Caudal)
                            caudal_str_match = re.search(r'[\d\.]+', parts[1])
                            if caudal_str_match:
                                caudal_m3s = float(caudal_str_match.group(0))

                    except (IndexError, ValueError) as e:
                        print(f"Advertencia: No se pudo procesar el valor de Nivel/Caudal '{nivel_caudal_str}'. Error: {e}")

                    datos_extraidos['datos_hidrometricos'].append({
                        'nombre_estacion': nombre_estacion,
                        'nivel_m': nivel_m,
                        'caudal_m3s': caudal_m3s
                    })
            print(f"Datos Hidrométricos estáticos extraídos: {len(datos_extraidos['datos_hidrometricos'])} registros.")
        else:
            print("No se encontró la tabla de Datos Hidrométricos (se esperaba en el índice 13).")

        # --- EXTRAER 11. ESTADO DE PASOS FRONTERIZOS ---
        datos_extraidos['estado_pasos_fronterizos'] = []
        if len(document.tables) > 17:
            table_pasos = document.tables[17]
            for i, row in enumerate(table_pasos.rows):
                if i == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 3 and any(cells):
                    datos_extraidos['estado_pasos_fronterizos'].append({
                        'nombre_paso': cells[0],
                        'condicion': cells[1],
                        'observaciones': cells[2]
                    })
            print(f"Estado de Pasos Fronterizos extraído: {len(datos_extraidos['estado_pasos_fronterizos'])} registros.")
        else:
            print("No se encontró la tabla de Estado de Pasos Fronterizos.")

        # --- GUARDAR LOS DATOS EXTRAÍDOS EN UN ARCHIVO JSON ---
        os.makedirs(DATA_OUTPUT_FOLDER, exist_ok=True)
        with open(DATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(datos_extraidos, f, ensure_ascii=False, indent=4)
        print(f"Datos extraídos guardados en: {DATA_FILE}")

        return DATA_FILE

    except Exception as e:
        print(f"Error al extraer datos del Word: {e}")
        return None

# --- EJECUCIÓN DEL SCRIPT ---
def main():
    # --- Verificación de variables de entorno ---
    if not EMAIL_PASSWORD:
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] ERROR: La variable de entorno GMAIL_APP_PASSWORD no está configurada. El script no puede continuar.")
        exit()

    parser = argparse.ArgumentParser(description="Descarga y procesa el último informe de monitoreo.")
    parser.add_argument('--force', action='store_true', help='Forzar el procesamiento incluso si el informe ya fue procesado.')
    args = parser.parse_args()

    if args.force:
        print("--- MODO FORZADO ACTIVADO ---")

    ruta_informe_descargado, subject_del_correo = descargar_ultimo_informe(
        EMAIL_USER, EMAIL_PASSWORD, IMAP_SERVER, DOWNLOAD_FOLDER
    )

    if not ruta_informe_descargado:
        print("FALLO: No se pudo descargar el informe para procesar.")
        return

    # --- Lógica de verificación ---
    nombre_archivo = os.path.basename(ruta_informe_descargado)
    id_informe_nuevo = obtener_id_del_informe(nombre_archivo)

    if not id_informe_nuevo:
        print(f"ADVERTENCIA: No se pudo generar un ID para el archivo '{nombre_archivo}'. Se procesará de todas formas.")
    else:
        id_existente = None
        try:
            with open(DATA_FILE, 'r', encoding='utf-8') as f:
                datos_existentes = json.load(f)
                id_existente = datos_existentes.get('id_informe_origen')
        except (FileNotFoundError, json.JSONDecodeError):
            pass # Si el archivo no existe o está corrupto, procedemos.

        if not args.force and id_informe_nuevo == id_existente:
            print(f"VERIFICACIÓN: El informe '{id_informe_nuevo}' ya ha sido procesado. Para volver a procesarlo, use --force.")
            print("Proceso omitido para evitar sobreescribir datos.")
            os.remove(ruta_informe_descargado) # Limpiamos el docx descargado
            return # Termina la ejecución

    # --- Si pasa la verificación (o es forzado), procesa ---
    json_ruta = extraer_datos_docx(ruta_informe_descargado, subject_del_correo, id_informe_nuevo)

    if json_ruta:
        print(f"Proceso de extracción completado exitosamente para: {json_ruta}")
        try:
            os.remove(ruta_informe_descargado)
            print(f"Archivo de informe '{nombre_archivo}' eliminado después de procesar.")            
        except OSError as e:
            print(f"Error al intentar eliminar el archivo del informe: {e}")
    else:
        print("FALLO: No se pudo extraer y guardar los datos del informe.")


if __name__ == "__main__":
    os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)
    os.makedirs(DATA_OUTPUT_FOLDER, exist_ok=True)
    main()
